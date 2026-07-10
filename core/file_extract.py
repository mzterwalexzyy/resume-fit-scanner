"""
Converts an uploaded resume file (PDF/DOCX/plain text, base64-encoded) into
plain text, so it can flow into the exact same analyze_resume_fit(resume_text,
job_description_text) contract as a pasted resume. Kept separate from
analyze.py so the core text-in/JSON-out function signature never changes --
this is purely a boundary-level convenience.
"""
import base64
import io

SUPPORTED_FILE_TYPES = {"pdf", "docx", "txt"}
MAX_FILE_BYTES = 5 * 1024 * 1024  # 5MB -- plenty for a resume, caps memory use


class FileExtractionError(ValueError):
    """Raised for any problem decoding or parsing an uploaded resume file."""


def _extract_pdf(data: bytes) -> str:
    from pypdf import PdfReader

    try:
        reader = PdfReader(io.BytesIO(data))
    except Exception as e:
        raise FileExtractionError(f"Could not read the PDF file: {e}") from e
    if reader.is_encrypted:
        raise FileExtractionError("This PDF is password-protected; please upload an unlocked copy.")
    pages_text = [page.extract_text() or "" for page in reader.pages]
    return "\n".join(pages_text)


def _extract_docx(data: bytes) -> str:
    from docx import Document

    try:
        doc = Document(io.BytesIO(data))
    except Exception as e:
        raise FileExtractionError(f"Could not read the DOCX file: {e}") from e

    parts = [p.text for p in doc.paragraphs]
    # doc.paragraphs alone silently skips text inside real Word tables --
    # exactly the layout pattern check_table_like_layout looks for, so
    # dropping it here would make that formatting check blind on DOCX input.
    for table in doc.tables:
        for row in table.rows:
            parts.append(" | ".join(cell.text for cell in row.cells))
    return "\n".join(parts)


def _extract_txt(data: bytes) -> str:
    try:
        return data.decode("utf-8")
    except UnicodeDecodeError as e:
        raise FileExtractionError(f"Could not decode the text file as UTF-8: {e}") from e


_EXTRACTORS = {
    "pdf": _extract_pdf,
    "docx": _extract_docx,
    "txt": _extract_txt,
}


def extract_text_from_file(file_base64: str, file_type: str) -> str:
    """Decode a base64-encoded resume file and return its plain text.

    Raises FileExtractionError (a ValueError) on any problem: unsupported
    type, oversized payload, bad base64, corrupt/encrypted file, or a file
    that decodes to no usable text.
    """
    file_type = (file_type or "").strip().lower().lstrip(".")
    if file_type not in SUPPORTED_FILE_TYPES:
        raise FileExtractionError(
            f"Unsupported file type '{file_type}'. Supported: "
            f"{', '.join(sorted(SUPPORTED_FILE_TYPES))}."
        )

    try:
        data = base64.b64decode(file_base64, validate=True)
    except Exception as e:
        raise FileExtractionError(f"Could not decode file content (invalid base64): {e}") from e

    if len(data) == 0:
        raise FileExtractionError("The uploaded file is empty.")
    if len(data) > MAX_FILE_BYTES:
        raise FileExtractionError(
            f"The uploaded file is too large ({len(data)} bytes, max {MAX_FILE_BYTES})."
        )

    text = _EXTRACTORS[file_type](data)
    if not text.strip():
        raise FileExtractionError(
            "No text could be extracted from this file -- it may be a scanned "
            "image without a text layer. Please paste the resume text directly instead."
        )
    return text
