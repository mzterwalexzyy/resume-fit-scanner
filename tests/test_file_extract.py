"""
Tests for core/file_extract.py (PDF/DOCX/TXT -> plain text) and its wiring
into the MCP tool in mcp_server/server.py. Generates synthetic PDF/DOCX
resumes on the fly with fpdf2/python-docx -- no binary fixture files checked
into the repo.

Run with:  py -m tests.test_file_extract     (from the project root)
"""
import base64
import io
import os
import sys

sys.path.insert(0, os.path.dirname(os.path.dirname(os.path.abspath(__file__))))

from core.file_extract import FileExtractionError, extract_text_from_file

SAMPLE_RESUME_LINES = [
    "Jordan Rivera",
    "Experience",
    "Backend Engineer, Nimbus Cloud Systems -- Jun 2022 - Present",
    "Built REST API services in Python using Django and PostgreSQL.",
    "Skills",
    "Python, Django, PostgreSQL, Docker, AWS",
]


def _assert(condition, message):
    if not condition:
        raise AssertionError(message)


def _make_pdf_base64() -> str:
    from fpdf import FPDF

    pdf = FPDF()
    pdf.add_page()
    pdf.set_font("Helvetica", size=12)
    for line in SAMPLE_RESUME_LINES:
        pdf.cell(0, 10, text=line, new_x="LMARGIN", new_y="NEXT")
    raw = bytes(pdf.output())
    return base64.b64encode(raw).decode("ascii")


def _make_docx_base64() -> str:
    from docx import Document

    doc = Document()
    for line in SAMPLE_RESUME_LINES:
        doc.add_paragraph(line)
    buf = io.BytesIO()
    doc.save(buf)
    return base64.b64encode(buf.getvalue()).decode("ascii")


def _make_docx_with_table_base64() -> str:
    from docx import Document

    doc = Document()
    doc.add_paragraph("Jamie Lee")
    table = doc.add_table(rows=1, cols=2)
    cells = table.rows[0].cells
    cells[0].text = "Skills"
    cells[1].text = "Kubernetes, Terraform"
    buf = io.BytesIO()
    doc.save(buf)
    return base64.b64encode(buf.getvalue()).decode("ascii")


def main():
    pdf_b64 = _make_pdf_base64()
    pdf_text = extract_text_from_file(pdf_b64, "pdf")
    _assert("Python" in pdf_text and "Django" in pdf_text,
            "PDF extraction should recover the resume's actual text content")

    docx_b64 = _make_docx_base64()
    docx_text = extract_text_from_file(docx_b64, "docx")
    _assert("Python" in docx_text and "PostgreSQL" in docx_text,
            "DOCX extraction should recover the resume's actual text content")

    table_docx_b64 = _make_docx_with_table_base64()
    table_docx_text = extract_text_from_file(table_docx_b64, "docx")
    _assert("Kubernetes" in table_docx_text and "Terraform" in table_docx_text,
            "DOCX extraction must not silently drop text inside real Word tables")

    txt_b64 = base64.b64encode(b"Plain text resume with Python and AWS experience.").decode("ascii")
    txt_text = extract_text_from_file(txt_b64, "txt")
    _assert("Python" in txt_text, "TXT extraction should return the decoded text as-is")

    try:
        extract_text_from_file(pdf_b64, "exe")
        raise AssertionError("unsupported file type should have raised FileExtractionError")
    except FileExtractionError:
        pass

    try:
        extract_text_from_file("not-valid-base64!!!", "txt")
        raise AssertionError("invalid base64 should have raised FileExtractionError")
    except FileExtractionError:
        pass

    try:
        extract_text_from_file(base64.b64encode(b"").decode("ascii"), "txt")
        raise AssertionError("empty file should have raised FileExtractionError")
    except FileExtractionError:
        pass

    oversized = base64.b64encode(b"x" * (6 * 1024 * 1024)).decode("ascii")
    try:
        extract_text_from_file(oversized, "txt")
        raise AssertionError("oversized file should have raised FileExtractionError")
    except FileExtractionError:
        pass

    # End-to-end: the actual MCP tool function, not just the extraction helper.
    from mcp_server.server import analyze_resume_fit as tool_fn

    jd = (
        "Backend Engineer -- Fictional Startup Inc.\n\n"
        "About the role\nWe are looking for a Backend Engineer to help build "
        "and scale our core API for this growing company.\n\n"
        "Requirements\n- Strong experience with Python and Django\n"
        "- Experience with PostgreSQL and relational database design\n"
        "- Experience with Docker and AWS\n"
        "- Comfortable with Git and Agile/Scrum workflows\n"
    )
    result = tool_fn(job_description_text=jd, resume_file_base64=pdf_b64, resume_file_type="pdf")
    _assert("rejected" not in result, f"file-based call should succeed, got: {result}")
    _assert(isinstance(result["fit_score"], int), "file-based call should produce a real fit_score")

    both_result = tool_fn(
        job_description_text=jd,
        resume_text="some pasted text",
        resume_file_base64=pdf_b64,
        resume_file_type="pdf",
    )
    _assert(both_result.get("rejected") is True,
            "providing both resume_text and a file should be rejected, not silently pick one")

    print("All file-extraction assertions passed.")


if __name__ == "__main__":
    main()
